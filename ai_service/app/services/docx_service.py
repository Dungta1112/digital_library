from dataclasses import dataclass
from io import BytesIO
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from app.services import pdf_service

CHUNK_SIZE = pdf_service.CHUNK_SIZE
CHUNK_OVERLAP = pdf_service.CHUNK_OVERLAP
Chunk = pdf_service.Chunk
ExtractResult = pdf_service.ExtractResult

def extract_chunks(docx_bytes: bytes) -> ExtractResult:
    doc = Document(BytesIO(docx_bytes))
    current_heading = ""
    current_paras = []
    para_count = 0
    para_start = 0
    chunks = []
    chunk_idx = 0

    def flush_section():
        nonlocal chunk_idx, para_start
        if not current_paras:
            return
        text = " ".join(current_paras)
        pieces = pdf_service.split_text(text, CHUNK_SIZE, CHUNK_OVERLAP)
        num_paras = len(current_paras)
        for piece in pieces:
            chunks.append(Chunk(
                page=para_start,
                chunk_index=chunk_idx,
                text=piece,
                section=current_heading,
                para_start=para_start,
                para_end=para_start + num_paras - 1,
            ))
            chunk_idx += 1
        current_paras.clear()

    for para in doc.paragraphs:
        para_count += 1
        text = para.text.strip()
        if not text:
            continue
        style = para.style
        if style and style.type == WD_STYLE_TYPE.PARAGRAPH and 'Heading' in style.name:
            flush_section()
            current_heading = f"{style.name}: {text}"
            para_start = para_count
        else:
            if not current_paras:
                para_start = para_count
                if not current_heading:
                    current_heading = "Nội dung"
            current_paras.append(text)

    flush_section()
    return ExtractResult(
        pages_total=para_count,
        pages_with_text=para_count,
        chunks=chunks,
        pages_ocred=0,
    )

def is_readable_docx(docx_bytes: bytes) -> bool:
    try:
        Document(BytesIO(docx_bytes))
        return True
    except Exception:
        return False
